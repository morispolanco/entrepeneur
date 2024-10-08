import streamlit as st
import requests
import json
import re
from docx import Document
from io import BytesIO
import plotly.graph_objs as go
import pandas as pd

# Set page configuration
st.set_page_config(page_title="Entrepreneurship Feasibility Analysis", page_icon="🌎", layout="wide")

# Function to create the information column
def create_info_column():
    st.markdown("""
    ## About this application

    This application provides a feasibility analysis for starting different types of industries, businesses, or services in any city worldwide. It allows users to evaluate the viability of entrepreneurship in specific sectors based on the selected location. The analysis includes interactive charts where applicable.

    ### How to use the application:

    1. Enter a city and country.
    2. Select an industry, business, or service sector, or describe your own entrepreneurship idea.
    3. Click on "Get feasibility analysis" to generate the analysis.
    4. Read the detailed information provided and interact with the charts.
    5. If desired, download a DOCX document with all the information.

    ### Author and update:
    **Moris Polanco**, August 30, 2024


    ---
    **Note:** This application uses artificial intelligence to generate detailed analyses based on available online data. Always verify the information with official sources for a more accurate and up-to-date analysis.
    """)

# Titles and Main Column
st.title("Entrepreneurship Feasibility Analysis")

# Create two columns
col1, col2 = st.columns([1, 2])

# Column 1 content
with col1:
    create_info_column()

# Column 2 content
with col2:
    TOGETHER_API_KEY = st.secrets.get("TOGETHER_API_KEY")
    SERPER_API_KEY = st.secrets.get("SERPER_API_KEY")

    # Types of industry, business, or service
    sectors = [
        "Manufacturing Industry", "Retail Trade", "Tourism", "Information Technology", 
        "Agriculture", "Transport and Logistics", "Private Education", "Healthcare", "Restaurants and Cafes",
        "Construction", "Financial Services", "Handicrafts", "Audiovisual Production", "Renewable Energy", 
        "Telecommunications", "Consulting Services", "Real Estate", "Fashion and Textiles"
    ]

    def search_information(query):
        url = "https://google.serper.dev/search"
        payload = json.dumps({
            "q": f"{query} city feasibility analysis statistics"
        })
        headers = {
            'X-API-KEY': SERPER_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()

    def generate_feasibility_analysis(city, country, sector, context):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Context: {context}\n\nCity: {city}\nCountry: {country}\nSector: {sector}\n\nProvide a detailed and extensive feasibility analysis on starting a business in the '{sector}' sector in {city}, {country}. The information should be accurate, complete, and based on real data. Include statistics, relevant numerical data, market analysis, entry barriers, and any additional information that may be of interest. Make sure to cover multiple aspects of business feasibility.\n\nWhere possible, include numerical data that can be used to create charts. For example, you could provide market size projections over the next 5 years, or a breakdown of market share by competitors.\n\nFeasibility analysis:",
            "max_tokens": 4096,
            "temperature": 0.2,
            "top_p": 0.9,
            "top_k": 50,
            "repetition_penalty": 1.1,
            "stop": ["City:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def extract_numerical_data(text):
        # Extract data for line chart (e.g., market size projections)
        line_data = re.findall(r'(\d{4}).*?(\d+(?:\.\d+)?)', text)
        
        # Extract data for pie chart (e.g., market share)
        pie_data = re.findall(r'(\w+(?:\s+\w+)?)\s*:\s*(\d+(?:\.\d+)?)\s*%', text)
        
        return line_data, pie_data

    def create_charts(line_data, pie_data):
        charts = []
        
        if line_data:
            df = pd.DataFrame(line_data, columns=['Year', 'Value'])
            df['Year'] = pd.to_numeric(df['Year'])
            df['Value'] = pd.to_numeric(df['Value'])
            df = df.sort_values('Year')
            
            fig = go.Figure(data=go.Scatter(x=df['Year'], y=df['Value'], mode='lines+markers'))
            fig.update_layout(title='Market Size Projection', xaxis_title='Year', yaxis_title='Market Size')
            charts.append(fig)
        
        if pie_data:
            labels = [item[0] for item in pie_data]
            values = [float(item[1]) for item in pie_data]
            
            fig = go.Figure(data=[go.Pie(labels=labels, values=values)])
            fig.update_layout(title='Market Share')
            charts.append(fig)
        
        return charts

    def create_docx(city, country, information, sources):
        doc = Document()
        doc.add_heading(f'Feasibility Analysis - {city}, {country}', 0)

        doc.add_heading('Location', level=1)
        doc.add_paragraph(f'{city}, {country}')

        for sector, data in information.items():
            doc.add_heading(f'{sector}', level=2)
            doc.add_paragraph(data)

        doc.add_heading('Sources', level=1)
        for source in sources:
            doc.add_paragraph(source, style='List Bullet')

        doc.add_paragraph('\nNote: This document was generated by an AI assistant. Verify the information with official sources for more accurate and up-to-date data.')

        return doc

    st.write("Enter a city and country:")
    city = st.text_input("City:")
    country = st.text_input("Country:")

    st.write("Select an industry, business, or service sector, or describe your entrepreneurship idea:")
    sector_options = sectors + ["Describe your own entrepreneurship idea"]
    selected_sector = st.selectbox("Sector", sector_options)

    if selected_sector == "Describe your own entrepreneurship idea":
        business_idea = st.text_area("Describe your entrepreneurship idea:")
        if not business_idea:
            st.warning("Please describe your entrepreneurship idea.")
            st.stop()
        selected_sector = "Custom idea: " + business_idea  # Concatenate user idea into a separate sector

    if st.button("Get feasibility analysis"):
        if city and country and selected_sector:
            with st.spinner("Searching for information and generating feasibility analysis..."):
                information, all_sources = {}, []

                # Search for relevant information
                search_results = search_information(f"{city} {country} {selected_sector}")
                context = "\n".join([item["snippet"] for item in search_results.get("organic", [])])
                sources = [item["link"] for item in search_results.get("organic", [])]

                # Generate feasibility analysis
                data = generate_feasibility_analysis(city, country, selected_sector, context)

                information[selected_sector] = data
                all_sources.extend(sources)

                # Display the information
                st.subheader(f"Feasibility analysis for entrepreneurship in the sector: {selected_sector}")
                st.markdown(f"**Location: {city}, {country}**")
                st.write(data)

                # Extract numerical data and create charts
                line_data, pie_data = extract_numerical_data(data)
                charts = create_charts(line_data, pie_data)

                # Display charts
                for chart in charts:
                    st.plotly_chart(chart)

                st.write("---")

                # Button to download the document
                doc = create_docx(city, country, information, all_sources)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Download analysis as DOCX",
                    data=buffer,
                    file_name=f"Feasibility_Analysis_{city.replace(' ', '_')}_{country.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Please enter a city and country, and select a sector or describe your entrepreneurship idea.")
